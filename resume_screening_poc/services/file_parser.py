import base64
from io import BytesIO

from services.llm_client import LLMConfig, call_vision_ocr

IMAGE_RESUME_EXTENSIONS = ("png", "jpg", "jpeg")
SUPPORTED_RESUME_EXTENSIONS = ("pdf", "docx", "txt", *IMAGE_RESUME_EXTENSIONS)
SUPPORTED_RESUME_EXTENSION_SET = set(SUPPORTED_RESUME_EXTENSIONS)
PDF_OCR_RENDER_SCALE = 1.5
UNSUPPORTED_RESUME_FORMAT_MESSAGE = "当前版本暂不支持该文件格式，请上传 PDF、DOCX、TXT、PNG、JPG 或 JPEG 文件。"


def get_resume_file_extension(file_name: str) -> str:
    return file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""


def is_supported_resume_file(uploaded_file) -> bool:
    return get_resume_file_extension(uploaded_file.name) in SUPPORTED_RESUME_EXTENSION_SET


def parse_resume_file(uploaded_file, ocr_config: LLMConfig | None = None) -> tuple[str, str | None]:
    """
    解析上传的简历文件。

    返回：
    - 文本内容
    - 错误信息；成功时为 None

    说明：
    可复制文字的 PDF/DOCX/TXT 会优先直接提取文本。
    图片文件和扫描版 PDF 会调用视觉模型做 OCR。
    """
    file_name = uploaded_file.name
    extension = get_resume_file_extension(file_name)

    if extension not in SUPPORTED_RESUME_EXTENSION_SET:
        return "", UNSUPPORTED_RESUME_FORMAT_MESSAGE

    file_bytes = uploaded_file.getvalue()

    try:
        if extension == "pdf":
            text = _parse_pdf(file_bytes)
            if not text.strip():
                text, ocr_error = _parse_pdf_with_ocr(file_bytes, ocr_config)
                if ocr_error:
                    return "", ocr_error
        elif extension == "docx":
            text = _parse_docx(file_bytes)
        elif extension in IMAGE_RESUME_EXTENSIONS:
            text, ocr_error = _parse_image_with_ocr(file_bytes, extension, ocr_config)
            if ocr_error:
                return "", ocr_error
        else:
            text = _parse_txt(file_bytes)
    except Exception as exc:
        return "", f"文件解析失败：{exc}"

    text = text.strip()
    if not text:
        return "", "未提取到文字内容。请换用更清晰的图片，或先人工转换为 TXT/DOCX 后再上传。"

    return text, None


def _parse_pdf(file_bytes: bytes) -> str:
    """使用 PyMuPDF 提取 PDF 文本。"""
    import fitz

    text_parts = []
    with fitz.open(stream=file_bytes, filetype="pdf") as document:
        for page in document:
            text_parts.append(page.get_text())
    return "\n".join(text_parts)


def _parse_pdf_with_ocr(file_bytes: bytes, ocr_config: LLMConfig | None) -> tuple[str, str | None]:
    """把扫描版 PDF 页面渲染成图片后交给视觉模型 OCR。"""
    images = _render_pdf_pages_for_ocr(file_bytes)
    return call_vision_ocr(images, ocr_config)


def _render_pdf_pages_for_ocr(file_bytes: bytes) -> list[dict[str, str]]:
    import fitz

    images = []
    with fitz.open(stream=file_bytes, filetype="pdf") as document:
        for page in document:
            pixmap = page.get_pixmap(matrix=fitz.Matrix(PDF_OCR_RENDER_SCALE, PDF_OCR_RENDER_SCALE), alpha=False)
            images.append(
                {
                    "mime_type": "image/png",
                    "base64": base64.b64encode(pixmap.tobytes("png")).decode("ascii"),
                }
            )
    return images


def _parse_image_with_ocr(
    file_bytes: bytes,
    extension: str,
    ocr_config: LLMConfig | None,
) -> tuple[str, str | None]:
    mime_type = "image/png" if extension == "png" else "image/jpeg"
    image = {
        "mime_type": mime_type,
        "base64": base64.b64encode(file_bytes).decode("ascii"),
    }
    return call_vision_ocr([image], ocr_config)


def _parse_docx(file_bytes: bytes) -> str:
    """使用 python-docx 提取 DOCX 段落和表格文本。"""
    from docx import Document

    document = Document(BytesIO(file_bytes))
    text_parts = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                text_parts.append(" | ".join(cells))

    return "\n".join(text_parts)


def _parse_txt(file_bytes: bytes) -> str:
    """优先按 UTF-8 读取 TXT；失败时尝试常见中文编码 GBK。"""
    for encoding in ("utf-8-sig", "utf-8", "gbk"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="ignore")
